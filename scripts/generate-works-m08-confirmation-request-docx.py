"""Generate client confirmation DOCX for Works M-08 decisions review."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Leave_Management"
    / "Works_M08_Client_Decisions_Confirmation_Request.docx"
)

# Better path - this is Works not Leave
OUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Works_M08_Client_Decisions_Confirmation_Request.docx"
)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                table.rows[r_i + 1].cells[c_i].text = str(val)
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "GAPLMB IOMS — Works / Construction (M-08)\n"
        "Client Decisions Review &\n"
        "Confirmation Request"
    )
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Based on: Must decide before build_updated.docx\n"
        "Purpose: Confirm accepted decisions, clarify ambiguous items, and approve proposed defaults\n"
        "Version: 1.0 | Status: Awaiting client confirmation"
    )
    run.italic = True

    h("1. Purpose")
    p(
        "We have reviewed the client inputs in “Must decide before build_updated.docx” for the "
        "Works / Work Order Management redevelopment (IOMS M-08)."
    )
    p(
        "Most decisions are clear and can be used for build. A small set of items are ambiguous or "
        "incomplete. This note lists: (a) decisions we have recorded as accepted, (b) items needing "
        "explicit confirmation, and (c) proposed defaults where unanswered — please tick Accept / Change."
    )
    p(
        "Note: This file relates to Works (M-08) only. Leave Management clarifications are a separate "
        "document and are not covered here."
    )

    h("2. Decisions Recorded as Accepted")
    p("Please confirm these match your intent. If any is wrong, mark Change and write the correct rule.")
    add_table(
        ["#", "Topic", "Our understanding of your decision", "Confirm"],
        [
            [
                "A1",
                "Payments vs M-06",
                "Every work payment shall be linked to M-06 Payment Voucher (expenditure).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A2",
                "Work Order No.",
                "User-entered (not auto-generated).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A3",
                "Existing works data",
                "Current DB rows are test data and may be deleted (no migration).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A4",
                "Mobilization advance recovery",
                "Adjustable against later bills; recovery amount is user-entered (not automatic); "
                "advance can reduce to zero.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A5",
                "Over-billing (bills + advance > WO value)",
                "Allow with DA override + remark (not hard block).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A6",
                "Advance cap base",
                "10% cap on WO Amount excluding GST (fallback: Tender Value).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A7",
                "Number of advances per WO",
                "One advance record per Work Order.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A8",
                "Multi-bill payment",
                "One payment can cover several unpaid bills.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A9",
                "Edit / delete of bills & payments",
                "Locked once payment mode and details are recorded (no edit/delete thereafter).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A10",
                "WO amendment after Approved",
                "Not required in v1 (values entered after approvals).",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A11",
                "Mark Work complete",
                "Manual “Mark completed” by DA; overdue if end date passed.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A12",
                "License / approvals table",
                "Deferred to v2.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A13",
                "Scope / T&C / DLP / penalty / retention",
                "Optional text/number on registration in v1; no auto retention deduction.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A14",
                "Vendor master",
                "One shared vendor master; AMC can use later.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A15",
                "Payment modes",
                "Cash / Cheque / DD / Online.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A16",
                "Dashboard date filter",
                "Filter by Work Order date.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A17",
                "GST on vendor bills",
                "Bills are GST-inclusive; capture GST %, GST amount, and bill/taxable amount; "
                "maintain bill-wise and work-wise GST totals.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A18",
                "TDS on Payment Voucher",
                "Capture TDS Applicable Y/N, Section, Rate %, Gross, TDS Amount, Net Payable; "
                "link to Vendor / Bill / Voucher / Work; support TDS reports.",
                "☐ Accept  ☐ Change: ___",
            ],
            [
                "A19",
                "Security Deposit / PBG",
                "Record SD/PBG amount and mode (Cash / DD / BG / Other) with mode-specific fields "
                "as listed in your inputs.",
                "☐ Accept  ☐ Change: ___",
            ],
        ],
    )

    h("3. Items Needing Explicit Confirmation (Ambiguous / Incomplete)")
    p(
        "Please answer each item below. Where helpful, a Proposed Default is shown — "
        "tick Accept Default or write your rule."
    )

    h("3.1 Approval depth (CRITICAL)", 2)
    p(
        "Your note says: “WO and every bill / advance / payment? Apply the DO → DV → DA.” "
        "This reads as Option B. Please tick one:"
    )
    bullet("☐ Option A — DO → DV → DA on Work Order only; bills / advance / payments are recording after WO is Approved")
    bullet("☐ Option B — DO → DV → DA on Work Order AND every bill / advance / payment")
    p("If Option B: confirm advance and payment documents each have their own Draft → Verified → Approved cycle. ☐ Yes ☐ No")

    h("3.2 GST field model", 2)
    p("Please confirm:")
    bullet("☐ Bill total = Taxable amount + GST amount (GST-inclusive total)")
    bullet("☐ GST stored as single GST amount in v1 (no CGST/SGST split)")
    bullet("☐ OR CGST + SGST split required in v1")
    bullet("☐ GST details entered on Bill (recommended)")
    bullet("☐ GST details entered on Payment only")
    bullet("☐ GST details on both Bill and Payment")
    p("Proposed default if unanswered: Taxable + GST% + GST ₹ + Total on Bill; single GST amount; no CGST/SGST split in v1.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.3 TDS base (“Applicable Amount”)", 2)
    p("Formula uses Applicable Amount × TDS Rate. Which base should apply by default?")
    bullet("☐ Taxable value excluding GST (recommended)")
    bullet("☐ Gross bill amount including GST")
    bullet("☐ User always types Applicable Amount manually")
    p("Proposed default: TDS on taxable value excl. GST; user may override Applicable Amount.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.4 When bills / payments become locked", 2)
    p("You said locked once payment mode and details are recorded. Please clarify:")
    bullet("☐ Lock as soon as any payment is saved (even Draft)")
    bullet("☐ Lock only after M-06 Payment Voucher is Approved (recommended)")
    p("Proposed default: Editable until first Approved payment voucher posts against the bill.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.5 Advance adjustment at payment time", 2)
    p("Please confirm recovery mechanics:")
    bullet("☐ At each bill payment, user may enter “Advance adjusted (₹)”")
    bullet("☐ Total adjustments cannot exceed approved advance amount")
    bullet("☐ Remaining advance = Approved advance − sum of adjustments (can reach zero)")
    p("Proposed default: all three above = Yes.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.6 Mobilization advance ceiling", 2)
    p("Mockups stated max 10% of Work Order value. Please confirm:")
    bullet("☐ Max approved advance = 10% of WO Amount excl. GST")
    bullet("☐ Other %: ______")
    p("Proposed default: 10% of WO Amount excl. GST.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.7 Security Deposit / PBG lifecycle (v1)", 2)
    bullet("☐ Allow multiple SD/PBG records per Work Order")
    bullet("☐ v1 = record only (no refund / invoke workflow)")
    bullet("☐ Refund / release workflow required in v1")
    bullet("☐ Cash/DD SD should create M-06 document in v1  ☐ Later  ☐ Never")
    p("Proposed default: Multiple records allowed; record-only in v1; M-06 link later.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("3.8 New scope in v1 (GST + TDS + SD/PBG)", 2)
    p("Your updated file adds GST, TDS, and SD/PBG as requirements. Please confirm v1 scope:")
    bullet("☐ Include GST bill-wise capture + work GST summary in Works v1")
    bullet("☐ Include TDS on M-06 Payment Voucher + TDS reports in Works v1")
    bullet("☐ Include SD/PBG register in Works v1")
    p("Proposed default: All three included in Works v1.")
    bullet("☐ Accept default  ☐ Change / defer: _________________")

    h("3.9 Vendor mandatory before Work Order", 2)
    bullet("☐ Vendor must be registered/selected before WO save (recommended)")
    bullet("☐ Free-text contractor name still allowed in v1")
    p("Proposed default: Vendor master selection mandatory.")
    bullet("☐ Accept default  ☐ Change: _________________")

    h("4. Proposed Defaults Summary (for quick approval)")
    p("If you prefer not to answer item-by-item, you may approve this whole default pack:")
    add_table(
        ["Item", "Proposed default"],
        [
            ["Approval depth", "NEEDS YOUR EXPLICIT A or B choice (no silent default)"],
            ["GST model", "On Bill: Taxable + GST% + GST ₹ + Total; single GST (no split)"],
            ["TDS base", "Taxable excl. GST; override allowed"],
            ["Bill lock", "After first Approved M-06 payment"],
            ["Advance recovery", "User-entered adjustment on payment; cannot exceed remaining advance"],
            ["Advance cap", "10% of WO Amount excl. GST"],
            ["SD/PBG", "Multiple records; record-only in v1"],
            ["GST + TDS + SD/PBG", "In Works v1"],
            ["Vendor", "Mandatory from vendor master"],
            ["Retention auto-deduct", "No in v1"],
        ],
    )
    bullet("☐ I approve the Proposed Defaults pack above (except Approval depth A/B which I tick in Section 3.1)")
    bullet("☐ I do not approve — see changes marked in Sections 2–3")

    h("5. Client Response Block")
    add_table(
        ["Field", "Details"],
        [
            ["Organisation", "GAPLMB"],
            ["Document responded", "Must decide before build_updated.docx"],
            ["Responded by (Name)", ""],
            ["Designation", ""],
            ["Date", ""],
            ["Signature", ""],
        ],
    )
    p("Overall decision:")
    bullet("☐ Proceed to build with accepted decisions + my answers in Section 3")
    bullet("☐ Proceed with Proposed Defaults pack + my A/B choice in 3.1")
    bullet("☐ Hold — further discussion required")
    p("Additional remarks: ________________________________________________________________")
    p("________________________________________________________________________________")

    h("6. Document Control")
    add_table(
        ["Field", "Value"],
        [
            ["Title", "Works M-08 Client Decisions Confirmation Request"],
            ["Version", "1.0"],
            ["Module", "IOMS M-08 Construction / Works"],
            ["Next step after reply", "Freeze scope and start Works workflow implementation"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
