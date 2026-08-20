"""Generate GAPLMB Leave Policy + Ready-to-Rollout Holiday Calendar DOCX."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Leave_Management"
    / "GAPLMB_Leave_Policy_and_Holiday_Calendar_2026.docx"
)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                table.rows[r_i + 1].cells[c_i].text = str(val)
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "THE GOA AGRICULTURAL PRODUCE &\n"
        "LIVESTOCK MARKETING BOARD (GAPLMB)\n\n"
        "LEAVE POLICY & RULES\n"
        "with\n"
        "HOLIDAY CALENDAR 2026\n"
        "(Ready to Rollout)"
    )
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "For Client Review / Adoption in IOMS (M-01 Leave)\n"
        "Status: DRAFT FOR APPROVAL — Proposed rules marked; open items in Part D\n"
        "Version: 1.0 | Calendar Year: 2026"
    )
    run.italic = True

    h("Document Structure")
    bullet("Part A — Leave Policy (General Principles & Process)")
    bullet("Part B — Leave Types & Rules to be Followed")
    bullet("Part C — Holiday Calendar 2026 (Ready to Rollout)")
    bullet("Part D — Missing / Open Rules Requiring Client Inputs")
    bullet("Part E — Adoption & Sign-off")

    # ========== PART A ==========
    h("PART A — LEAVE POLICY (GENERAL PRINCIPLES & PROCESS)")

    h("A1. Purpose & Scope", 2)
    p(
        "This Leave Policy governs sanction, recording, and accounting of leave for employees "
        "of the Goa Agricultural Produce & Livestock Marketing Board (GAPLMB), and defines "
        "the holiday calendar used for leave calculations in the IOMS Leave module (M-01)."
    )
    p(
        "Unless GAPLMB issues a local circular to the contrary, leave entitlements and "
        "conditions follow Central Civil Services (Leave) Rules / Swamy FR & SR Part III "
        "practice as reflected in operational documents supplied by the Board."
    )
    p("Applies to: Regular GAPLMB staff on IOMS employee master (unless excluded in Part D).")
    p("Does not by itself cover: Daily-rated / casual workers except where Part C Annexure-IV note applies (client confirmation required).")

    h("A2. Guiding Principles", 2)
    bullet("Leave cannot be claimed as a matter of right. Sanctioning authority may refuse or revoke leave of any kind.")
    bullet("Authority shall not alter the kind of leave applied for by the employee (except conversion rules where separately allowed).")
    bullet("Earned Leave should not ordinarily be denied in the last ten years of service.")
    bullet("Unauthorised absence may be regularised as Extraordinary Leave only by competent authority.")
    bullet("Continuous leave of any kind shall not ordinarily exceed five years (except with higher sanction / disability provisions).")
    bullet("Employee on leave shall not take up other employment without prior sanction.")
    bullet("All leave in IOMS shall follow workflow: Application (DO) → Verification (DV) → Approval / Rejection (DA).")

    h("A3. Roles & Responsibilities", 2)
    add_table(
        ["Role", "Responsibility"],
        [
            ["Employee / DO", "Submit correct application, dates, grounds, documents, substitute (where required)"],
            ["Controlling / Reporting Officer (DV)", "Verify facts, recommend / return with remarks"],
            ["Secretary / DA", "Sanction or reject; Order issuance; ensure balance debit"],
            ["HR / Admin", "Maintain holiday calendar, opening balances, file numbering, service-book entry checklist"],
            ["Accounts", "Receive Order copy for pay / leave salary implications"],
        ],
    )

    h("A4. Application & Sanction Process (Rules to Follow)", 2)
    bullet("Leave must be applied in IOMS before proceeding on leave, except ex-post facto cases permitted under policy.")
    bullet("CL / RH / Special Holiday: short application path (letter-equivalent fields).")
    bullet("EL / HPL / Commuted / ML / PL (and other Form-1 types): Form-1 equivalent fields including prefix/suffix, address during leave, LTC flag, substitute.")
    bullet("DV verifies; DA approves or rejects with reason code and remarks where required.")
    bullet("On Approval: leave balance is debited (as per type rules) and Sanction Order PDF is generated.")
    bullet("Employee shall not leave headquarters during leave without permission stated on Order (where applicable).")
    bullet("On return: employee resumes same post/station unless Order states otherwise; medical fitness certificate where required.")

    h("A5. Prefix / Suffix / Weekly Offs", 2)
    bullet("Proposed weekly offs: Saturday and Sunday (all locations unless Part D exception list is approved).")
    bullet("Prefix / suffix of Saturday, Sunday and Public Holidays may be permitted as shown on Sanction Order.")
    bullet("Proposed debit rule: prefix/suffix days are NOT debited from leave balance; only the sanctioned leave days are debited (confirm in Part D).")
    bullet("Public Holidays falling inside an EL/CCL spell count as leave days for that spell (per CCS practice for EL/CCL).")

    h("A6. Ex-Post Facto Leave", 2)
    bullet("Ex-post facto sanction may be accorded where leave was availed and later regularised by competent authority.")
    bullet("IOMS shall mark such cases clearly; Order wording shall include “Ex-post facto sanction…”.")
    bullet("Who may create ex-post facto applications: to be confirmed in Part D (proposed: Admin / authorised HR only).")

    h("A7. Overlap, Cancellation & Rejection", 2)
    bullet("No overlapping leave requests (Pending / Verified / Approved) for the same employee.")
    bullet("Rejected / Cancelled leave does not debit balance.")
    bullet("Cancellation after Approval: only with DA/Admin process and balance restoration (confirm in Part D).")

    # ========== PART B ==========
    h("PART B — LEAVE TYPES & RULES TO BE FOLLOWED")

    h("B1. Leave Type Catalogue (Proposed for IOMS v1)", 2)
    add_table(
        ["Code", "Leave Type", "Debited From", "Typical Max / Notes"],
        [
            ["EL", "Earned Leave", "EL account", "Cap 300 days; max 180 days per spell (exceptions in rules)"],
            ["HPL", "Half Pay Leave", "HPL account", "With or without MC as rules allow"],
            ["COMMUTED", "Commuted Leave", "HPL @ 2× days", "Usually on MC; not for LPR"],
            ["CL", "Casual Leave", "CL account", "8 days/year; max 3 consecutive days; half-day allowed"],
            ["RH", "Restricted Holiday", "RH account", "Max 2 per calendar year from RH list"],
            ["SPL_H", "Special Holiday", "SPL.H / duty-in-lieu", "Per Special Holiday annexure / compensatory rules"],
            ["ML", "Maternity Leave", "Not debited", "Typically 180 days"],
            ["PL", "Paternity Leave", "Not debited", "15 days; eligibility conditions"],
            ["CCL", "Child Care Leave", "Separate CCL account", "Optional v1 / Phase E — see Part D"],
            ["EOL", "Extraordinary Leave", "No debit", "Optional v1 / Phase E — see Part D"],
        ],
    )

    h("B2. Earned Leave (EL)", 2)
    bullet("Credit: 15 days in advance on 1 January and 15 days on 1 July each year.")
    bullet("Accumulation: maximum 300 days (excluding LTC encashment days where separately allowed).")
    bullet("When balance is between 286 and 300 at credit time: advance credit held in set-off and adjusted as per CCS practice (Excel pattern 300(+n)).")
    bullet("Mid-year appointment / retirement / resignation: credit @ 2½ days per completed calendar month (as applicable).")
    bullet("Reduction: credit reduced by 1/10th of EOL / dies-non in previous half-year (subject to max of available credit).")
    bullet("Rounding: fractions rounded to nearest day.")
    bullet("Maximum availed at a time: ordinarily 180 days (higher limits only as per rules for specified cases / LPR).")
    bullet("Application: Form-1; Order mandatory on sanction.")

    h("B3. Half Pay Leave (HPL)", 2)
    bullet("CCS practice: 10 days on 1 January and 10 days on 1 July (confirm vs current IOMS config of 20 on 1 Jan — Part D).")
    bullet("May be granted with or without medical certificate as per rules.")
    bullet("Overstayal without sanction may be debited to HPL then EOL (disciplinary implications).")

    h("B4. Commuted Leave", 2)
    bullet("Granted on medical certificate generally; not exceeding half of HPL due.")
    bullet("Debit: twice the number of days of Commuted Leave from HPL account.")
    bullet("May be taken even if EL is available.")
    bullet("Not to be used as leave preparatory to retirement.")
    bullet("If employee quits without returning: recalculation / recovery as per rules (death/incapacitation exceptions).")
    bullet("Application: Form-1 + MC; Order on sanction.")

    h("B5. Casual Leave (CL)", 2)
    bullet("Annual credit: 8 days per calendar year (as per GAPLMB balance register practice).")
    bullet("Maximum 3 consecutive days per application (IOMS enforced).")
    bullet("Half-day CL permitted (as reflected in staff balance sheet).")
    bullet("CL is not combined with Paternity Leave (CCS note).")
    bullet("Application: short letter format (CL template).")

    h("B6. Restricted Holiday (RH)", 2)
    bullet("Employee may avail any two Restricted Holidays from Annexure-III list in the calendar year.")
    bullet("RH date must match an entry on the Restricted Holiday list for that year.")
    bullet("Application: short RH letter format; occasion / festival to be stated.")
    bullet("Balance tracked as RH remaining (opening typically 2).")

    h("B7. Special Holiday (SPL.H)", 2)
    bullet("Governed by Annexure-II Special Holidays and GAPLMB compensatory practice.")
    bullet("Offices remain functional on Special Holidays unless they fall on Saturday/Sunday (GAD notification).")
    bullet("Compensatory format (GAPLMB letter): Special Holiday on <date> in lieu of duty attended on <date>.")
    bullet("Special Holidays may be prefixed or suffixed to other leave (GAD note).")
    bullet("Where Special Holiday falls on Sunday: request for holiday or day-in-lieu not allowed (2026 Mahashivratri note).")
    bullet("Milad-Un-Nabi (2026): if not availed on notified day, may be availed on any other working day in 2026 only (GAD note) — confirm Board adoption in Part D.")

    h("B8. Maternity Leave (ML)", 2)
    bullet("Female employee: typically 180 days (as per GAPLMB sanctioned Orders).")
    bullet("Not debited to leave account.")
    bullet("Supporting documents mandatory.")
    bullet("Form-1 application; Sanction Order on approval.")
    bullet("Special Maternity Leave (stillbirth / death of child soon after birth): 60 days — confirm inclusion in v1 (Part D).")
    bullet("Not admissible for threatened abortion; admissible for induced abortion (CCS).")

    h("B9. Paternity Leave (PL)", 2)
    bullet("Male employee with less than two surviving children: 15 days.")
    bullet("Avail during confinement of wife: from 15 days before delivery up to 6 months after delivery; or on adoption / pre-adoption foster care of child below one year (within 6 months).")
    bullet("Leave salary equal to last pay drawn; not debited to leave account.")
    bullet("May be combined with any leave except Casual Leave.")
    bullet("Not to be refused normally; lapses if not applied within allowed window.")
    bullet("Form-1; supporting documents as required.")

    h("B10. Child Care Leave (CCL) — Proposed Optional / Phase E", 2)
    bullet("Female employees and single male employees (unmarried / widower / divorcee) with minor children.")
    bullet("Maximum 730 days in entire service for up to two eldest surviving children.")
    bullet("Max 3 spells/year (6 for single female); each spell minimum 5 days.")
    bullet("First 365 days: 100% pay; next 365 days: 80% pay.")
    bullet("Not debited to EL/HPL account; maintained separately.")
    bullet("Include in v1? — Client decision in Part D.")

    h("B11. Extraordinary Leave (EOL) — Proposed Optional / Phase E", 2)
    bullet("May be granted when no other leave is admissible, or when other leave is admissible but applied for EOL.")
    bullet("Limits differ for temporary vs permanent servants; continuous leave ceiling of five years applies.")
    bullet("EOL may regularise absence without leave retrospectively.")
    bullet("Include in v1? — Client decision in Part D.")

    h("B12. Encashment (EL)", 2)
    bullet("IOMS already supports retirement encashment calculator: (Basic Pay + DA) / 30 × EL days.")
    bullet("Actual payment process remains with Accounts / existing finance procedure (confirm Part D).")

    # ========== PART C ==========
    doc.add_page_break()
    h("PART C — HOLIDAY CALENDAR 2026 (READY TO ROLLOUT)")

    banner = doc.add_paragraph()
    run = banner.add_run(
        "READY TO ROLLOUT — Calendar Year 2026\n"
        "Source: Government of Goa, General Administration Department Notification "
        "No.37/8/2025-GAD-III/4171 dated 30/09/2025\n"
        "Adoption: Proposed for all GAPLMB offices / yards / HO unless Part D states otherwise."
    )
    run.bold = True

    h("C1. How to Use This Calendar in Leave Processing", 2)
    add_table(
        ["Category", "Office Status", "Leave System Use"],
        [
            ["Public Holiday (Annexure I)", "Closed", "Prefix/suffix; sandwich; block overlapping leave as public holiday"],
            ["Special Holiday (Annexure II)", "Normally open*", "SPL.H rules; may prefix/suffix other leave"],
            ["Restricted Holiday (Annexure III)", "Open", "Employee may take any 2 as RH leave"],
            ["Weekly off (Sat–Sun)", "Off (proposed)", "Prefix/suffix"],
            ["Commercial & Industrial (Annexure IV)", "See Part D", "Not applied to regular staff unless confirmed"],
            ["Bank Holidays (Annexure V)", "N/A for leave", "Informational only — not used for staff leave"],
        ],
    )
    p("* Unless the Special Holiday falls on Saturday/Sunday.")
    p("* Dates marked with asterisk (*) are subject to appearance of the moon — Admin must update IOMS when confirmed.")

    h("C2. Annexure I — Public Holidays 2026 (Office Closed)", 2)
    add_table(
        ["Sr", "Holiday", "Date", "Day"],
        [
            ["1", "Republic Day", "26 Jan 2026", "Monday"],
            ["2", "Holi (Dhulivandana)", "03 Mar 2026", "Tuesday"],
            ["3", "Gudi Padava", "19 Mar 2026", "Thursday"],
            ["4", "Id-Ul Fitr (Ramzan Id)*", "21 Mar 2026", "Saturday"],
            ["5", "Ram Navami", "26 Mar 2026", "Thursday"],
            ["6", "Good Friday", "03 Apr 2026", "Friday"],
            ["7", "Birth Anniversary of Dr. Babasaheb Ambedkar", "14 Apr 2026", "Tuesday"],
            ["8", "May Day", "01 May 2026", "Friday"],
            ["9", "Id-Ul-Zuha (Bakri Id)*", "28 May 2026", "Thursday"],
            ["10", "Independence Day", "15 Aug 2026", "Saturday"],
            ["11", "Hartalika / Ganesh Chaturthi (1st Day)", "14 Sep 2026", "Monday"],
            ["12", "Ganesh Chaturthi (2nd Day)", "15 Sep 2026", "Tuesday"],
            ["13", "Gandhi Jayanti", "02 Oct 2026", "Friday"],
            ["14", "Dussehra (Vijaya Dashmi)", "20 Oct 2026", "Tuesday"],
            ["15", "Diwali (Deepavali)", "08 Nov 2026", "Sunday"],
            ["16", "Feast of St. Francis Xavier", "03 Dec 2026", "Thursday"],
            ["17", "Goa Liberation Day", "19 Dec 2026", "Saturday"],
            ["18", "Christmas Day", "25 Dec 2026", "Friday"],
        ],
    )

    h("C3. Annexure II — Special Holidays 2026", 2)
    add_table(
        ["Sr", "Holiday", "Date", "Day", "Notes"],
        [
            [
                "1",
                "Mahashivratri",
                "15 Feb 2026",
                "Sunday",
                "Falls on Sunday — no holiday / day-in-lieu request allowed",
            ],
            [
                "2",
                "Milad-Un-Nabi / Id-e-Milad*",
                "25 Aug 2026",
                "Tuesday",
                "Must be sanctioned; if not availed on date, may be availed any other working day in 2026 only",
            ],
        ],
    )
    p("N.B.: Special Holidays can be prefixed or suffixed to any other kind of leave.")

    h("C4. Annexure III — Restricted Holidays 2026 (Choose Any Two)", 2)
    add_table(
        ["Sr", "Holiday", "Date", "Day"],
        [
            ["1", "New Year Day", "01 Jan 2026", "Thursday"],
            ["2", "Makarsankranti", "14 Jan 2026", "Wednesday"],
            ["3", "Feast of St. Joseph Vaz", "16 Jan 2026", "Friday"],
            ["4", "Guru Ravi Das Birthday", "01 Feb 2026", "Sunday"],
            ["5", "Shivaji Jayanti", "19 Feb 2026", "Thursday"],
            ["6", "Mahavir Jayanti", "31 Mar 2026", "Tuesday"],
            ["7", "Maundy Thursday", "02 Apr 2026", "Thursday"],
            ["8", "Vaisakhi / Vishu", "14 Apr 2026", "Tuesday"],
            ["9", "Budha Purnima", "01 May 2026", "Friday"],
            ["10", "Feast of Sacred Heart of Jesus", "12 Jun 2026", "Friday"],
            ["11", "Muharam", "26 Jun 2026", "Friday"],
            ["12", "Onam", "26 Aug 2026", "Wednesday"],
            ["13", "Raksha Bandhan", "28 Aug 2026", "Friday"],
            ["14", "Janmashtami", "04 Sep 2026", "Friday"],
            ["15", "All Souls Day", "02 Nov 2026", "Monday"],
            ["16", "Govardhan Puja", "10 Nov 2026", "Tuesday"],
            ["17", "Bhaubij", "11 Nov 2026", "Wednesday"],
            ["18", "Guru Nanak's Birthday", "24 Nov 2026", "Tuesday"],
            ["19", "Guru Teg Bahadur Martyrdom Day", "24 Nov 2026", "Tuesday"],
            ["20", "Feast of Immaculate Conception of Mary", "08 Dec 2026", "Tuesday"],
            ["21", "Christmas Eve", "24 Dec 2026", "Thursday"],
            ["22", "New Year's Eve", "31 Dec 2026", "Thursday"],
        ],
    )
    p("Rule: Each employee may avail only TWO Restricted Holidays from this list during calendar year 2026.")

    h("C5. Annexure IV — Commercial & Industrial Holidays 2026 (Reference)", 2)
    p(
        "Listed for reference. Apply to GAPLMB regular staff only if client confirms in Part D "
        "(proposed default: Do NOT apply to regular staff leave calendar)."
    )
    add_table(
        ["Sr", "Holiday", "Date", "Day"],
        [
            ["1", "Republic Day", "26 Jan 2026", "Monday"],
            ["2", "Birth Anniversary of Dr. Babasaheb Ambedkar", "14 Apr 2026", "Tuesday"],
            ["3", "May Day", "01 May 2026", "Friday"],
            ["4", "Independence Day", "15 Aug 2026", "Saturday"],
            ["5", "Hartalika / Ganesh Chaturthi (1st Day)", "14 Sep 2026", "Monday"],
            ["6", "Gandhi Jayanti", "02 Oct 2026", "Friday"],
            ["7", "Diwali (Deepavali)", "08 Nov 2026", "Sunday"],
            ["8", "Goa Liberation Day", "19 Dec 2026", "Saturday"],
            ["9", "Christmas Day", "25 Dec 2026", "Friday"],
        ],
    )
    p(
        "Note (GOI MoF memo): Casual employees including daily-rated staff entitled to paid holidays "
        "if in service on preceding and succeeding working days — confirm applicability to GAPLMB."
    )

    h("C6. Weekly Offs (Proposed Ready to Rollout Default)", 2)
    add_table(
        ["Day", "Status", "Remark"],
        [
            ["Saturday", "Weekly off", "Used for prefix/suffix"],
            ["Sunday", "Weekly off", "Used for prefix/suffix"],
            ["Monday–Friday", "Working days", "Unless Public Holiday / sanctioned leave"],
        ],
    )

    h("C7. IOMS Load Checklist (Operations)", 2)
    bullet("☐ Load Annexure I as category=Public for year 2026")
    bullet("☐ Load Annexure II as category=Special for year 2026")
    bullet("☐ Load Annexure III as category=Restricted for year 2026")
    bullet("☐ Mark moon-dependent dates as Tentative")
    bullet("☐ Configure weekly offs Sat+Sun (or approved exception matrix)")
    bullet("☐ Publish calendar to staff / notice board / IOMS help text")
    bullet("☐ Assign Admin owner for mid-year moon confirmation updates")

    # ========== PART D ==========
    doc.add_page_break()
    h("PART D — MISSING / OPEN RULES REQUIRING CLIENT INPUTS")
    p(
        "The items below are incomplete, ambiguous, or not evidenced in the documents folder. "
        "Please answer each item. Until answered, IOMS will use the Proposed Default shown."
    )

    h("D1. Policy Gaps (Missing Rules)", 2)
    add_table(
        ["#", "Missing / Open Rule", "Proposed Default", "Client Decision / Notes"],
        [
            [
                "D1.1",
                "Exact legal basis adopted by GAPLMB (CCS Leave Rules as-is vs Board-specific Standing Orders)",
                "CCS/Swamy practice + GAD holidays + Board Orders",
                "",
            ],
            [
                "D1.2",
                "Staff categories covered (regular / contract / daily-rated / deputation)",
                "Regular staff on IOMS employee master only",
                "",
            ],
            [
                "D1.3",
                "HPL credit schedule",
                "Align to CCS 10+10 (Jan/Jul) — change from current 20 on 1 Jan",
                "",
            ],
            [
                "D1.4",
                "Meaning of Excel “Commuted Leave Balance”",
                "Treat as HPL account balance",
                "",
            ],
            [
                "D1.5",
                "Whether prefix/suffix days debit leave balance",
                "No debit for prefix/suffix; debit sanctioned leave days only",
                "",
            ],
            [
                "D1.6",
                "Sandwich rule when leave sits between holidays/weekends",
                "Follow CCS EL practice; confirm for CL/RH",
                "",
            ],
            [
                "D1.7",
                "CL combination with other leave",
                "CL not combinable with PL; other combinations as CCS",
                "",
            ],
            [
                "D1.8",
                "RH on a Restricted Holiday that is also Sunday",
                "Allow RH only if date is on list and is a working day otherwise? Confirm",
                "",
            ],
            [
                "D1.9",
                "Adoption of movable Milad rule for Board staff",
                "Adopt GAD Annexure-II note as-is",
                "",
            ],
            [
                "D1.10",
                "Include CCL in v1",
                "No — Phase E",
                "",
            ],
            [
                "D1.11",
                "Include EOL / Leave Not Due / leave conversion in v1",
                "No — Phase E",
                "",
            ],
            [
                "D1.12",
                "Special Maternity Leave (60 days) in v1",
                "Yes if ML in scope; else document later",
                "",
            ],
            [
                "D1.13",
                "Child Adoption Leave (180 days) in v1",
                "Phase E unless client marks Yes",
                "",
            ],
            [
                "D1.14",
                "Cancellation after Approval & balance restore",
                "Allowed only by DA/Admin with audit",
                "",
            ],
            [
                "D1.15",
                "Overstayal auto-debit HPL/EOL",
                "Manual regularisation in v1; auto later",
                "",
            ],
            [
                "D1.16",
                "Fitness certificate gate before return from medical leave",
                "Warn in v1; hard block later",
                "",
            ],
            [
                "D1.17",
                "Checkpost / shift weekly offs different from Sat–Sun",
                "None — same as HO unless exception list given",
                "",
            ],
            [
                "D1.18",
                "Apply Annexure IV / V to any staff",
                "No",
                "",
            ],
            [
                "D1.19",
                "Substitute mandatory before Approve",
                "Yes for EL/ML/PL/Commuted; optional for CL/RH",
                "",
            ],
            [
                "D1.20",
                "Self-service apply by employee login in v1",
                "Yes if employee has login; else HR/DO on behalf",
                "",
            ],
            [
                "D1.21",
                "Order file numbering formula",
                "GAPLMB/{file}/ADM-{YYYY}/{seq} — client to confirm file codes",
                "",
            ],
            [
                "D1.22",
                "Balance certificate “as on” date on Order",
                "Next half-year end (30 Jun / 31 Dec) as in samples",
                "",
            ],
            [
                "D1.23",
                "Encashment payment workflow integration",
                "Calculator only in v1; payment offline",
                "",
            ],
            [
                "D1.24",
                "Service Book auto-entry vs checklist only",
                "Checklist on Order in v1",
                "",
            ],
            [
                "D1.25",
                "Language of printed forms (English only vs bilingual)",
                "English only",
                "",
            ],
        ],
    )

    h("D2. Missing Document Formats (Attach or Confirm N/A)", 2)
    add_table(
        ["#", "Format Needed", "Why Needed", "Client Action"],
        [
            ["F1", "Service Book leave entry specimen", "Digital checklist / later auto-entry", "☐ Attach / ☐ N/A"],
            ["F2", "Medical Certificate format accepted by Board", "Commuted / HPL / ML validations", "☐ Attach / ☐ N/A"],
            ["F3", "Fitness Certificate format", "Return to duty after medical leave", "☐ Attach / ☐ N/A"],
            ["F4", "Paternity Leave Sanction Order sample", "Order PDF parity", "☐ Attach / ☐ N/A"],
            ["F5", "CCL application / Order (if in scope)", "CCL module", "☐ Attach / ☐ Out of scope"],
            ["F6", "Child Adoption Leave formats", "If adopted in policy", "☐ Attach / ☐ Out of scope"],
            ["F7", "EOL / LND application format", "Phase E", "☐ Attach / ☐ Out of scope"],
            ["F8", "Leave conversion request format (within 30 days)", "Phase E", "☐ Attach / ☐ Out of scope"],
            ["F9", "Rejection / return memo format", "DV/DA communications", "☐ Attach / ☐ Free text OK"],
            ["F10", "Board Standing Orders / local leave circulars", "Override CCS where applicable", "☐ Attach / ☐ None"],
            ["F11", "Letterhead / logo / seal for Order PDF", "Print-ready Orders", "☐ Attach"],
            ["F12", "Approver matrix (if not always Secretary)", "Workflow routing", "☐ Attach / ☐ Secretary only"],
            ["F13", "Exception weekly-off matrix by location", "Prefix/suffix engine", "☐ Attach / ☐ No exceptions"],
            ["F14", "Contract / daily-rated leave rules", "If covered by policy", "☐ Attach / ☐ Out of scope"],
        ],
    )

    h("D3. Data Inputs Before Go-Live", 2)
    add_table(
        ["#", "Input", "Format", "Due / Owner"],
        [
            ["1", "Emp ID mapping for all staff in balance Excel", "Excel", ""],
            ["2", "Opening balances as on cut-over date", "Excel", ""],
            ["3", "Cut-over date", "DD/MM/YYYY", ""],
            ["4", "Reporting officer for each employee", "Excel / IOMS update", ""],
            ["5", "Order signatory name & designation", "Text", ""],
            ["6", "ADM file-code list used in Orders", "List", ""],
            ["7", "Confirm Softcopy_Leave_Order.docx as print master", "Yes/No + file", ""],
            ["8", "Moon-date update owner (name)", "Text", ""],
        ],
    )

    h("D4. Client Quick Response Grid", 2)
    add_table(
        ["Ref", "Question", "Answer (Yes/No/Value)"],
        [
            ["D1.3", "HPL = 10+10 Jan/Jul?", ""],
            ["D1.4", "Commuted Excel column = HPL?", ""],
            ["D1.5", "Prefix/suffix not debited?", ""],
            ["D1.10", "CCL in v1?", ""],
            ["D1.11", "EOL in v1?", ""],
            ["D1.17", "Any weekly-off exceptions?", ""],
            ["D1.18", "Apply Annexure IV?", ""],
            ["D1.19", "Substitute mandatory?", ""],
            ["D1.21", "File number formula confirmed?", ""],
            ["C", "Adopt 2026 calendar Parts C2–C4 as official?", ""],
            ["—", "Other remarks", ""],
        ],
    )

    # ========== PART E ==========
    h("PART E — ADOPTION & SIGN-OFF")
    p("We approve adoption of this Leave Policy (Parts A–B) and Holiday Calendar 2026 (Part C) for IOMS rollout, subject to Part D answers attached.")
    add_table(
        ["Role", "Name", "Signature", "Date"],
        [
            ["Secretary / Competent Authority", "", "", ""],
            ["HR / Administration", "", "", ""],
            ["Accounts (acknowledged)", "", "", ""],
            ["IOMS Project Owner", "", "", ""],
        ],
    )
    p("Decision:")
    bullet("☐ Adopt as Ready to Rollout (with Part D defaults where unanswered)")
    bullet("☐ Adopt with Part D answers attached (mandatory for listed items)")
    bullet("☐ Revise and resubmit")

    h("Document Control")
    add_table(
        ["Field", "Value"],
        [
            ["Title", "GAPLMB Leave Policy & Rules with Holiday Calendar 2026 (Ready to Rollout)"],
            ["Version", "1.0"],
            ["Related proposal", "GAPLMB_Leave_Module_Redevelopment_Proposal.docx"],
            ["Source folder", "docs/Leave_Management"],
            ["Next step", "Complete Part D → Approve Part E → Configure IOMS calendar & leave types"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
