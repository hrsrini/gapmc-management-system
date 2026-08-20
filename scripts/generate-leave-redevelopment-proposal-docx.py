"""Generate client-review DOCX for Leave Management redevelopment proposal."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "docs" / "Leave_Management" / "GAPLMB_Leave_Module_Redevelopment_Proposal.docx"


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                table.rows[r_i + 1].cells[c_i].text = str(val)
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GAPLMB IOMS — Leave Management Module\nRedevelopment Proposal")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("For Client Review & Formal Approval")
    run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Module: M-01 HR / Leave\n"
        "Organisation: Goa Agricultural Produce & Livestock Marketing Board (GAPLMB)\n"
        "Basis: Existing IOMS leave screens + documents in docs/Leave_Management\n"
        "Document status: Proposal (implementation pending approval)\n"
        "Version: 1.0"
    )

    h("1. Purpose")
    p(
        "This document seeks formal client approval to redevelop the existing Leave Management "
        "capability in IOMS from a basic leave request and balance screen into a workflow-enabled "
        "leave office process aligned with GAPLMB operational practice, CCS/Swamy leave rules "
        "(as applied in practice), the Goa holiday calendar, and existing paper formats "
        "(Form-1 applications, CL/RH/Special Holiday letters, and Sanction Orders)."
    )
    p(
        "After approval of scope and after client inputs listed in Section 19 are received, "
        "implementation will proceed in agreed phases."
    )

    h("2. Executive Summary")
    add_table(
        ["Item", "Proposal"],
        [
            ["Approach", "Extend existing M-01 Leave module (do not build a separate application)"],
            ["Outcome", "Digital apply → verify → sanction → balance debit → Order PDF → service-book note"],
            ["Retain", "DO → DV → DA workflow, employee master, audit log, M-01 permissions"],
            [
                "Add",
                "Typed leave catalogue, holiday master, prefix/suffix, RH/SPL.H/Commuted/ML/PL rules, "
                "Sanction Order PDF, substitute duty",
            ],
            [
                "Out of scope (v1)",
                "Full EOL/LND edge cases, payroll auto-posting, UIDAI; Market Fee Returns form (unrelated document)",
            ],
            ["Approval requested", "Phases A–D (foundation through Sanction Order). Phase E optional later"],
        ],
    )

    h("3. Current System (As Built)")
    h("3.1 Screens", 2)
    bullet("Leave requests: /hr/leaves")
    bullet("Leave balances: /hr/leave-balances")

    h("3.2 Already Working", 2)
    bullet("Create / edit leave request in Pending status")
    bullet("Workflow: Pending → Verified → Approved / Rejected; DV return to Pending; Admin retrospective flag")
    bullet("Casual Leave (CL) maximum 3 consecutive calendar days per application")
    bullet("Overlap check with other active leave; holiday overlap via configuration JSON")
    bullet("Supporting document required for Maternity Leave (ML) / Child Care Leave (CCL)")
    bullet("Reporting-officer check for DV verification")
    bullet("Balance debit on Approve when a balance row exists")
    bullet(
        "Accrual cron stubs: EL 15 days on 1 Jan & 1 Jul; CL 8 days on 1 Jan; HPL 20 days on 1 Jan; EL cap warning at 300"
    )
    bullet("EL encashment calculator: (Basic + DA) / 30 × EL days")

    h("3.3 Gaps Versus Client Documents", 2)
    bullet("No structured holiday calendar (Public / Special / Restricted categories)")
    bullet("Leave types incomplete versus GAPLMB practice (RH, Special Holiday, Commuted, Paternity under-specified)")
    bullet("No prefix / suffix calculation used on Sanction Orders")
    bullet("Form-1 fields missing (address during leave, LTC flag, substitute, leave HQ, controlling officer remarks)")
    bullet("No Sanction Order PDF or official file numbering")
    bullet("Half-day CL not supported")
    bullet("Commuted leave not debited as 2 × days from Half Pay Leave (HPL)")
    bullet("EL set-off near 300 days (Excel pattern 300(+n)) not modelled")
    bullet("Opening balances from staff Excel not formally imported")

    h("4. Source Documents Reviewed")
    add_table(
        ["Document", "Use in Redevelopment"],
        [
            ["Leave_Rules.pdf", "Accrual, caps, Commuted / HPL / EL / ML / PL rules"],
            ["Holiday_list.pdf (Goa GAD 2026)", "Public / Special / Restricted holiday master"],
            ["Leave_Balance_of_Staff.xlsx", "Opening balances and leave-type reality check"],
            ["CL / RH / Special Holiday application formats", "Short application UX"],
            ["Form-1 + Maternity / Paternity / Commuted samples", "Full application UX"],
            ["EL / Checkpost / Ex-post facto / Softcopy Order", "Sanction Order PDF template"],
            ["Market_Fee_Returns_Format.pdf", "NOT leave-related — exclude from this project"],
        ],
    )

    h("5. Target Process (To-Be)")
    p("Employee / DO applies")
    p("        ↓")
    p("DV (Controlling / Reporting Officer) verifies + remarks")
    p("        ↓")
    p("DA (Secretary / authorised Approver) approves or rejects")
    p("        ↓")
    p("On Approve: debit balance + generate Sanction ORDER PDF")
    p("        ↓")
    p("Notify / copy: Employee, Substitute, Section/Yard, Accounts, Personal file")
    p("        ↓")
    p('Optional: "Entered on Service Book" checklist')

    h("5.1 Paper Role Mapping", 2)
    add_table(
        ["Paper Role", "IOMS Role"],
        [
            ["Applicant", "Employee / Data Originator (DO)"],
            ['Controlling Officer ("may be granted")', "Data Verifier (DV)"],
            ["Secretary (ORDER)", "Data Approver (DA)"],
        ],
    )

    h("6. Functional Scope — Leave Types (v1)")
    add_table(
        ["Code", "Name", "Typical Debit", "Key Rules"],
        [
            [
                "EL",
                "Earned Leave",
                "1:1 from EL",
                "Accrual 15 on Jan 1 & Jul 1; cap 300; max 180 days/spell (rules); prefix/suffix; Order PDF",
            ],
            ["HPL", "Half Pay Leave", "1:1 from HPL", "Accrual per approved schedule; MC as applicable"],
            ["COMMUTED", "Commuted Leave", "2 × days from HPL", "Medical/grounds; Form-1; Order PDF"],
            ["CL", "Casual Leave", "1:1 (allow 0.5)", "Annual 8; max 3 consecutive days; short letter format"],
            ["RH", "Restricted Holiday", "1 day from RH", "Max 2 per calendar year; date must be on Restricted list"],
            [
                "SPL_H",
                "Special Holiday",
                "Per special-holiday rules",
                "Compensatory (duty date) and/or Goa Special Holiday annexure rules",
            ],
            ["ML", "Maternity Leave", "No balance debit", "Typically 180 days; supporting docs mandatory; Form-1 + Order"],
            ["PL", "Paternity Leave", "No balance debit", "Typically 15 days; eligibility (less than 2 surviving children)"],
        ],
    )
    p(
        'Client confirmation required: Excel column "Commuted Leave Balance" appears to represent the HPL account '
        "(large balances). System will store HPL and display Commuted available as approximately half of HPL "
        "(floor of HPL divided by 2) in the UI."
    )
    p(
        "Deferred to Phase E (optional): EOL limits, Leave Not Due, leave-type conversion within 30 days, "
        "fitness certificate return-to-duty automation, vacation-department accrual."
    )

    h("7. Holiday Calendar (v1)")
    h("7.1 Categories (from Goa GAD notification for 2026)", 2)
    bullet("Public holidays (Annexure I) — office closed; used for prefix/suffix and sandwich logic")
    bullet(
        "Special holidays (Annexure II) — office normally open; Special Holiday rules "
        "(e.g. Milad movable; no lieu if falls on Sunday)"
    )
    bullet("Restricted holidays (Annexure III) — employee may avail any two in the year")
    p("Weekly offs for prefix/suffix (proposed default): Saturday + Sunday. Client to confirm checkpost or yard exceptions.")

    h("7.2 System Behaviour", 2)
    bullet("Admin maintains holidays by calendar year (seed 2026 from Annexures I–III)")
    bullet("Moon-dependent dates flagged as tentative for Admin update")
    bullet("Leave apply validates RH against Restricted list")
    bullet("Prefix/suffix auto-computed from weekends + Public holidays")

    h("8. Application Experience")
    h("8.1 Short Applications (CL / RH / SPL.H)", 2)
    p("Digital equivalent of existing letter formats:")
    bullet("Days / date(s)")
    bullet("Reason / occasion")
    bullet("Show live balances: CL, RH, SPL.H")
    bullet('For compensatory Special Holiday: "Special Holiday on ___ in lieu of duty attended on ___"')

    h("8.2 Form-1 Applications (EL / HPL / Commuted / ML / PL)", 2)
    p("Capture or auto-fill:")
    bullet("Name, post, office/section (from employee master)")
    bullet("Pay (from employee master where available)")
    bullet("Nature and period of leave")
    bullet("Prefix / suffix (auto-calculated; editable with audit if needed)")
    bullet("Grounds")
    bullet("Last leave return (auto from history)")
    bullet("LTC proposed Yes / No")
    bullet("Address during leave")
    bullet("Supporting documents (mandatory for ML and medical Commuted as configured)")
    bullet("Substitute employee(s)")
    bullet("Leave headquarters / destination (where applicable)")
    bullet("Controlling officer remarks (at DV step)")

    h("8.3 Ex-post Facto", 2)
    bullet("Authorised path using existing retrospective flag (Admin / policy)")
    bullet('ORDER wording: "Ex-post facto sanction is hereby accorded…"')

    h("9. Workflow, Statuses and Actions")
    add_table(
        ["Status", "Who Acts", "Actions"],
        [
            ["Pending", "DV / Admin", "Verify, or return with remarks"],
            ["Verified", "DA / Admin", "Approve, or Reject (reason code + remarks)"],
            ["Approved", "—", "Terminal; balance debited; ORDER generated"],
            ["Rejected / Cancelled", "—", "Terminal; no debit"],
        ],
    )
    p("Rules retained: DO/DV/DA segregation; no self-approve; reporting-officer DV routing.")

    h("10. Balance and Accrual Engine (v1)")
    add_table(
        ["Topic", "Behaviour"],
        [
            ["Opening balances", "Import from approved Excel after employee matching sign-off"],
            ["On Approve", "Debit debitDays (prefix/suffix handled per agreed rule)"],
            ["Commuted", "Debit 2 × leave days from HPL"],
            ["CL half-day", "Supported (0.5)"],
            ["RH", "Entitlement 2/year; track availed"],
            ["EL near 300", "Set-off credit logic per rules (Excel pattern 300(+n))"],
            ["Accrual job", "Keep daily cron; align HPL schedule after client confirmation"],
            ["Insufficient balance", "Hard block on Approve"],
        ],
    )

    h("11. Sanction Order PDF (v1)")
    p("Generate on Approve, based on Softcopy_Leave_Order.docx and sample Orders:")
    bullet("Header: Office of GAPLMB")
    bullet("File no.: GAPLMB/{ADM-file}/ADM-{YYYY}/{seq}")
    bullet("READ: application reference and date")
    bullet("Sanction paragraph: type, days, from–to, prefix/suffix, leave HQ")
    bullet("Continuance / reposting clauses")
    bullet("Balance certificate as on a stated date")
    bullet("Signatory: Secretary (configurable)")
    bullet("Copy to: employee, substitute, section/yard, Accounts, Personal file, Guard file")
    bullet("Checklist: Entered on Service Book")

    h("12. Data / Technical Outline")
    p("Extend (not replace):")
    bullet(
        "leave_requests — prefix/suffix, debit days, substitute, address, LTC flag, leave HQ, "
        "file no., order PDF URL, ex-post facto flag"
    )
    bullet("employee_leave_balances — retain; ensure types EL / HPL / CL / RH / SPL_H")
    bullet("New: hr_holidays (year, date, name, category, tentative)")
    bullet("New: leave order sequence / file numbering")
    bullet("Config: weekly offs, RH entitlement, ML/PL defaults, Order signatory title")
    p("UI: enhance Leave Requests and Leave Balances; Admin holiday calendar; Order download on approved request.")

    h("13. Phased Delivery")
    add_table(
        ["Phase", "Deliverable", "Outcome"],
        [
            ["A", "Holiday master + 2026 seed; leave-type catalogue; balance import tool", "Foundation"],
            ["B", "Type-aware apply UX (short + Form-1); prefix/suffix; RH/SPL.H; substitute", "Usable digital applications"],
            [
                "C",
                "Balance engine fixes (Commuted 2×, half-day CL, EL set-off, hard insufficient balance)",
                "Correct leave accounts",
            ],
            ["D", "Sanction Order PDF + file no. + copy list + service-book checklist", "Paper parity for Orders"],
            [
                "E (optional)",
                "EOL / LND / conversion / fitness return / advanced accrual edge cases",
                "Full CCS depth",
            ],
        ],
    )
    p("Recommended approval: Phases A–D as Leave Redevelopment v1.0.")

    h("14. What Will Not Change Without Separate Approval")
    bullet("IOMS login, roles (DO / DV / DA / ADMIN), M-01 permissions model")
    bullet("Employee master location (leave continues to use existing employees)")
    bullet("LTC / TA-DA / Tour as separate M-01 sub-features (already partially built)")
    bullet("Market fee returns digitisation (out of this proposal)")

    h("15. Preliminary Client Decisions (Summary)")
    p("Detailed input forms are in Section 19. Key decisions:")
    bullet("Leave types in v1: EL, HPL, Commuted, CL, RH, SPL.H, ML, PL — confirm or amend")
    bullet('Excel "Commuted Leave Balance" = HPL account — confirm')
    bullet("Weekly offs: Sat+Sun for all locations — confirm / list exceptions")
    bullet("DA for leave Orders: always Secretary — confirm / yard-wise")
    bullet("Approve blocked until substitute selected — Yes / No")
    bullet("HPL accrual: 20 once on 1 Jan (current) OR 10+10 on 1 Jan & 1 Jul (CCS)")
    bullet("Phases approved: A–D / A–E / custom")

    h("16. Acceptance Criteria (v1)")
    bullet("2026 Public / Special / Restricted holidays loadable and used in validations")
    bullet("Staff can apply CL / RH / SPL.H with correct entitlements")
    bullet("Form-1 types support prefix/suffix, address, substitute, documents")
    bullet("DV / DA workflow produces Approved leave with correct balance debit")
    bullet("Commuted leave reduces HPL by 2 × days")
    bullet("Sanction Order PDF matches Softcopy template fields")
    bullet("Opening balances importable from client Excel after mapping sign-off")
    bullet("Ex-post facto Orders supported for authorised retrospective cases")

    h("17. Risks and Dependencies")
    bullet("Employee name matching for Excel import (manual mapping sheet may be required)")
    bullet("Moon-dependent holiday date changes mid-year (Admin update process)")
    bullet("CCS handbook versus GAPLMB local practice — client confirms where local practice prevails")
    bullet("Service Book integration depth depends on existing service-book feature maturity")

    h("18. Approval Block")
    add_table(
        ["Role", "Name", "Signature", "Date"],
        [
            ["Client / Business Owner", "", "", ""],
            ["GAPLMB Admin / Secretary nominee", "", "", ""],
            ["Project / IOMS Owner", "", "", ""],
        ],
    )
    p("Decision:")
    bullet("☐ Approved Phases A–D")
    bullet("☐ Approved with changes (attach notes)")
    bullet("☐ Deferred")
    p("Approved scope notes: _______________________________________________________________")

    # Section 19 — Client inputs
    doc.add_page_break()
    h("19. Client Inputs Required (Clarifications, Information, Formats & Attachments)")
    p(
        "This section lists everything needed from the client before and during implementation. "
        "Please complete the tables, tick options, and attach the requested formats/samples. "
        "Incomplete items may delay the corresponding phase."
    )

    h("19.1 Scope & Policy Clarifications (Please Tick / Fill)", 2)

    p("A. Leave types to include in v1.0")
    bullet("☐ EL (Earned Leave)")
    bullet("☐ HPL (Half Pay Leave)")
    bullet("☐ Commuted Leave")
    bullet("☐ CL (Casual Leave)")
    bullet("☐ RH (Restricted Holiday)")
    bullet("☐ SPL.H (Special Holiday)")
    bullet("☐ ML (Maternity Leave) — default days: ______ (proposed 180)")
    bullet("☐ PL (Paternity Leave) — default days: ______ (proposed 15)")
    bullet("☐ EOL (Extraordinary Leave) — include in v1? ☐ Yes ☐ No (proposed: Phase E)")
    bullet("☐ Other (specify): _________________________________")

    p("B. Balance interpretation")
    bullet('☐ Confirm Excel "Commuted Leave Balance" = HPL account days')
    bullet("☐ Or provide exact meaning / formula of that column: _________________")
    bullet("☐ Confirm RH entitlement = 2 per calendar year")
    bullet("☐ Confirm CL entitlement = 8 per calendar year (half-days allowed? ☐ Yes ☐ No)")
    bullet("☐ Confirm EL accumulation cap = 300 days (with set-off credit near cap? ☐ Yes ☐ No)")

    p("C. Accrual schedule")
    bullet("☐ EL: 15 days on 1 Jan and 15 days on 1 Jul (proposed)")
    bullet("☐ HPL: ☐ 20 days once on 1 Jan (current system)  OR  ☐ 10 days on 1 Jan + 10 on 1 Jul (CCS)")
    bullet("☐ CL: 8 days credited on ☐ 1 Jan  OR  ☐ other date: ________")
    bullet("☐ Mid-year joiners: apply pro-rata credit? ☐ Yes ☐ No")

    p("D. Weekly offs & locations")
    bullet("☐ Saturday + Sunday are weekly offs for all yards / HO / checkposts")
    bullet("☐ Exceptions (list location + offs): _________________________________")
    bullet("☐ Sandwich rule (leave spanning weekends/public holidays): describe expected behaviour: ________")

    p("E. Authority & workflow")
    bullet("☐ DV = Reporting / Controlling Officer of employee (proposed)")
    bullet("☐ DA / Order signatory = Secretary for all leave (proposed)")
    bullet("☐ Or yard-wise / designation-wise approver matrix (please attach)")
    bullet("☐ Self-application by employee login required in v1? ☐ Yes ☐ No (HR applies on behalf)")
    bullet("☐ Approve must require substitute employee? ☐ Yes ☐ No")
    bullet("☐ Ex-post facto: who may create? ☐ Admin only  ☐ Admin + HR  ☐ Other: ________")

    p("F. Debit vs absence period")
    bullet(
        "☐ Prefix/suffix weekends & public holidays are NOT debited from leave balance "
        "(only shown on Order) — proposed"
    )
    bullet("☐ Or debit includes prefix/suffix — specify: ________")
    bullet("☐ Special Holiday / RH: always 1 day debit — confirm")

    p("G. Documents & medical")
    bullet("Mandatory supporting documents for: ☐ ML  ☐ PL  ☐ Commuted (MC)  ☐ HPL (MC)  ☐ Other: ________")
    bullet("Fitness certificate required before return to duty after medical leave? ☐ Yes (v1) ☐ Later ☐ No")
    bullet("Max upload size / allowed formats for leave docs: ________ (proposed PDF/JPG/PNG, 10 MB)")

    h("19.2 Information / Master Data to Provide", 2)
    add_table(
        ["#", "Information Required", "Format Preferred", "Owner / Due"],
        [
            ["1", "Final employee list mapped to IOMS Emp ID (for balance import)", "Excel", ""],
            [
                "2",
                "Opening leave balances as on go-live date (EL, HPL/Commuted, CL, RH, SPL.H)",
                "Excel (update Leave_Balance_of_Staff or new sheet)",
                "",
            ],
            ["3", "Go-live / cut-over date for balances", "Date (DD/MM/YYYY)", ""],
            ["4", "Reporting officer for each employee (if not already in IOMS)", "Excel", ""],
            ["5", "Designation → leave eligibility matrix (if any differences by cadre)", "Excel / PDF", ""],
            ["6", "Authorised Order signatory name & designation text", "Text", ""],
            ["7", "ADM file number series currently used (e.g. GAPLMB/386/ADM-2026/…)", "List / samples", ""],
            ["8", "Standard Copy-to distribution list by leave type / location", "Word / Excel", ""],
            [
                "9",
                "List of yards / checkposts / sections for posting & substitute pool",
                "Excel (or confirm IOMS yards sufficient)",
                "",
            ],
            ["10", "Holiday calendar for current + next year if beyond 2026 GAD list", "PDF / Excel", ""],
            ["11", "Local GAPLMB circulars that override CCS rules (if any)", "PDF", ""],
            ["12", "Contact for Accounts Section notification after sanction", "Name / email", ""],
        ],
    )

    h("19.3 Document Formats & Samples Still Needed (Please Attach)", 2)
    p(
        "The following formats/samples will be used to finalise screens and PDF layouts. "
        "Tick if already provided in docs/Leave_Management; otherwise attach latest signed specimen."
    )
    add_table(
        ["#", "Format / Sample", "Status", "Client Action"],
        [
            ["1", "Form-1 Leave Application (blank + filled EL)", "Provided (docx + samples)", "☐ Confirm latest version"],
            ["2", "Casual Leave application letter", "Provided (JPEG)", "☐ Confirm latest version"],
            ["3", "Restricted Holiday application letter", "Provided (JPEG)", "☐ Confirm latest version"],
            ["4", "Special Holiday application letter", "Provided (JPEG)", "☐ Confirm latest version"],
            ["5", "Commuted Leave Form-1 + Order", "Provided (PDF)", "☐ Confirm latest version"],
            ["6", "Maternity Leave Form-1 + Order", "Provided (PDF)", "☐ Confirm latest version"],
            ["7", "Paternity Leave Form-1 (+ Order if any)", "Form provided; Order?", "☐ Attach PL Sanction Order sample if used"],
            [
                "8",
                "Earned Leave Sanction Order (multi-day + single-day)",
                "Provided",
                "☐ Confirm Softcopy Word is the print master",
            ],
            ["9", "Ex-post facto Sanction Order", "Provided", "☐ Confirm when wording is mandatory"],
            [
                "10",
                "Checkpost / outstation leave Order variant",
                "Provided (JPG)",
                "☐ Confirm if template differs from HO",
            ],
            ["11", "Service Book leave entry page / format", "Not in folder", "☐ Attach specimen page"],
            ["12", "Medical Certificate / Fitness Certificate formats accepted", "Not in folder", "☐ Attach accepted formats"],
            ["13", "Leave ledger / balance card format (if maintained)", "Excel provided", "☐ Confirm if paper ledger still required"],
            ["14", "Rejection / return memo format (if any)", "Not in folder", "☐ Attach or confirm free-text remarks OK"],
            [
                "15",
                "Notification / email / dak format for Order circulation",
                "Not in folder",
                "☐ Attach or confirm PDF + copy list enough",
            ],
            [
                "16",
                "Special Maternity / Child Care Leave formats (if in scope)",
                "Not provided",
                "☐ Attach or mark Out of Scope",
            ],
            ["17", "Any bilingual (Konkani/Marathi) leave forms required", "Unknown", "☐ Yes (attach) / ☐ English only"],
        ],
    )

    h("19.4 Holiday Calendar Inputs", 2)
    bullet("☐ Confirm 2026 GAD Annexures I–III as the official calendar for GAPLMB staff")
    bullet("☐ Provide process for moon-dependent date confirmation (Id / Milad) — who updates IOMS and by when")
    bullet('☐ Confirm Special Holiday rules for GAPLMB (especially movable Milad and "no lieu on Sunday")')
    bullet("☐ Provide 2027 (or next year) list when notified, for Admin load")
    bullet(
        "☐ Confirm whether Commercial/Industrial/Bank holiday annexures apply to any GAPLMB staff category "
        "(proposed: No)"
    )

    h("19.5 Order PDF Layout Inputs", 2)
    bullet("☐ Softcopy_Leave_Order.docx is the master print layout — ☐ Yes  ☐ No (attach revised Word)")
    bullet("☐ Letterhead / logo / seal image for PDF (file): _________________")
    bullet("☐ Signatory block text (name + designation): _________________")
    bullet("☐ Digital signature required in v1? ☐ Yes ☐ No (wet-ink / stamp after print)")
    bullet("☐ File numbering rule (please write formula): GAPLMB/_____/ADM-YYYY/_____")
    bullet("☐ Separate number series per leave type / section? ☐ Yes ☐ No")
    bullet("☐ Balance certificate as-on date rule: ☐ Half-year end  ☐ Order date  ☐ Other: ________")

    h("19.6 Opening Balance Import Checklist", 2)
    bullet(
        "☐ Provide balances as on cut-over date (not mixed 30.04.2026 / 30.06.2026 snapshots without reconciliation)"
    )
    bullet("☐ Clarify EL values shown as 300(+n) — confirm n is set-off credit to honour")
    bullet("☐ Mark employees who left / joined after the Excel date")
    bullet("☐ Provide Emp ID mapping sheet (Name | Designation | IOMS Emp ID | Active Y/N)")
    bullet("☐ Sign-off that imported balances are authoritative for go-live")

    h("19.7 Non-Leave Document in Folder", 2)
    p(
        "Market_Fee_Returns_Format.pdf appears to be a Market Fee Returns form (M-04 related), not leave. "
        "Please confirm: ☐ Move out of Leave_Management  ☐ Ignore for this project  "
        "☐ Actually required for leave (explain): ________"
    )

    h("19.8 Client Response Sheet (Fill and Return)", 2)
    add_table(
        ["Ref", "Question", "Client Answer"],
        [
            ["15/A", "v1 leave types list confirmed?", ""],
            ["15/B", "Commuted column = HPL?", ""],
            ["15/C", "HPL accrual schedule chosen", ""],
            ["15/D", "Weekly offs / exceptions", ""],
            ["15/E", "DA / signatory rule", ""],
            ["15/F", "Substitute mandatory on Approve?", ""],
            ["15/G", "Phases approved (A–D / A–E / other)", ""],
            ["19.3", "Missing formats attached? (list)", ""],
            ["19.5", "Order Word master confirmed?", ""],
            ["19.6", "Balance cut-over date", ""],
            ["—", "Other remarks", ""],
        ],
    )

    h("19.9 Submission Instructions", 2)
    bullet("Return this section completed (PDF/Word) with attachments to the IOMS project owner.")
    bullet("Preferred attachment naming: Leave_ClientInput_<topic>_<YYYYMMDD>.pdf")
    bullet(
        "Implementation of Phase A will start after: (1) Section 18 approval, and "
        "(2) minimum inputs 19.1 A–E, 19.2 items 1–3 and 6–7, 19.5 Softcopy confirmation."
    )

    h("20. Document Control")
    add_table(
        ["Field", "Value"],
        [
            ["Document title", "GAPLMB IOMS Leave Management Module Redevelopment Proposal"],
            ["Version", "1.0"],
            ["Prepared for", "Client review and approval"],
            ["Related folder", "docs/Leave_Management"],
            ["Next step after approval", "Phase A implementation (holiday master, leave types, balance import)"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
